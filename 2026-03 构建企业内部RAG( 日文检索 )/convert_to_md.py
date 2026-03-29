import os
import pandas as pd
from docx import Document
from pathlib import Path

def docx_to_md(docx_path, output_path):
    """将 Word 转换为带标题层级和表格的 Markdown"""
    doc = Document(docx_path)
    md_content = []
    
    for obj in doc.element.body:
        # 处理段落
        if obj.tag.endswith('p'):
            para = [p for p in doc.paragraphs if p._element == obj][0]
            text = para.text.strip()
            if not text: continue
            
            # 根据大纲级别转换标题 (Heading 1 -> #, Heading 2 -> ##)
            if para.style.name.startswith('Heading'):
                level = para.style.name.split()[-1]
                try:
                    md_content.append(f"{'#' * int(level)} {text}")
                except:
                    md_content.append(f"### {text}")
            else:
                md_content.append(text)
        
        # 处理表格
        elif obj.tag.endswith('tbl'):
            table = [t for t in doc.tables if t._element == obj][0]
            data = []
            for row in table.rows:
                data.append([cell.text.strip().replace('\n', ' ') for cell in row.cells])
            
            if data:
                df = pd.DataFrame(data[1:], columns=data[0])
                md_content.append(df.to_markdown(index=False))
        
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n\n'.join(md_content))

def excel_to_md(xlsx_path, output_path):
    """将 Excel 的每个 Sheet 转换为 Markdown 块"""
    excel_file = pd.ExcelFile(xlsx_path)
    md_content = []
    
    for sheet_name in excel_file.sheet_names:
        df = pd.read_excel(xlsx_path, sheet_name=sheet_name)
        # 清洗空行空列
        df = df.dropna(how='all').dropna(axis=1, how='all')
        
        if not df.empty:
            md_content.append(f"## Sheet: {sheet_name}")
            md_content.append(df.to_markdown(index=False))
            
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n\n'.join(md_content))

def batch_convert(input_dir):
    output_dir = Path(input_dir) / "converted_md"
    output_dir.mkdir(exist_ok=True)
    
    for file in os.listdir(input_dir):
        input_path = Path(input_dir) / file
        ext = input_path.suffix.lower()
        output_path = output_dir / f"{input_path.stem}.md"
        
        try:
            if ext == '.docx':
                docx_to_md(input_path, output_path)
                print(f"Success: {file} -> Markdown")
            elif ext == '.xlsx':
                excel_to_md(input_path, output_path)
                print(f"Success: {file} -> Markdown")
        except Exception as e:
            print(f"Error converting {file}: {e}")

if __name__ == "__main__":
    # 获取脚本当前所在的目录
    current_dir = os.path.dirname(os.path.abspath(__file__))
    print(f"正在处理目录: {current_dir}")
    batch_convert(current_dir)