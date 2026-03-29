import os
import pandas as pd
from docx import Document
from pathlib import Path
import sys

# 解决打印日文到控制台可能乱码的问题
sys.stdout.reconfigure(encoding='utf-8')

def docx_to_md(docx_path, output_path):
    doc = Document(docx_path)
    md_content = [f"# Source: {docx_path.name}"]
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text: continue
        if para.style.name.startswith('Heading'):
            level = para.style.name.split()[-1]
            md_content.append(f"{'#' * (int(level) if level.isdigit() else 3)} {text}")
        else:
            md_content.append(text)
    
    # 简单处理表格
    for table in doc.tables:
        data = [[cell.text.strip().replace('\n', ' ') for cell in row.cells] for row in table.rows]
        if data:
            df = pd.DataFrame(data[1:], columns=data[0])
            md_content.append(df.to_markdown(index=False))
            
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n\n'.join(md_content))

def excel_to_md(xlsx_path, output_path):
    excel_file = pd.ExcelFile(xlsx_path)
    md_content = [f"# Source: {xlsx_path.name}"]
    for sheet_name in excel_file.sheet_names:
        df = pd.read_excel(xlsx_path, sheet_name=sheet_name)
        df = df.dropna(how='all').dropna(axis=1, how='all')
        if not df.empty:
            md_content.append(f"## Sheet: {sheet_name}")
            md_content.append(df.to_markdown(index=False))
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n\n'.join(md_content))

def batch_convert():
    # 1. 确定路径
    base_dir = Path(__file__).parent.absolute()
    output_dir = base_dir / "converted_md"
    output_dir.mkdir(exist_ok=True)
    
    print(f"--- 正在扫描目录: {base_dir} ---")
    
    # 2. 递归搜索所有文件
    # rglob('*') 会深入所有子文件夹
    count = 0
    for file_path in base_dir.rglob('*'):
        # 跳过临时文件 (如 ~$Word.docx) 和输出目录自身
        if file_path.name.startswith('~$') or output_dir in file_path.parents:
            continue
            
        ext = file_path.suffix.lower()
        output_file = output_dir / f"{file_path.stem}_{file_path.parent.name}.md"
        
        if ext == '.docx':
            print(f"发现 Word: {file_path.name}")
            docx_to_md(file_path, output_file)
            count += 1
        elif ext in ['.xlsx', '.xls']:
            print(f"发现 Excel: {file_path.name}")
            excel_to_md(file_path, output_file)
            count += 1
            
    print(f"--- 处理完成，共转换 {count} 个文件 ---")
    if count == 0:
        print("提示：未发现 .docx 或 .xlsx 文件，请检查文件后缀或路径。")

if __name__ == "__main__":
    batch_convert()