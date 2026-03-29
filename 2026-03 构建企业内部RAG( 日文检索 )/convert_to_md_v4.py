import os
import pandas as pd
from docx import Document
from pathlib import Path
import sys
import pdfplumber
import win32com.client as win32  # 仅限 Windows 运行

# 解决打印日文到控制台可能乱码的问题
sys.stdout.reconfigure(encoding='utf-8')

def doc_to_docx(doc_path):
    """通过调用本地 Word 将 .doc 转为临时 .docx"""
    word = win32.gencache.EnsureDispatch('Word.Application')
    doc = word.Documents.Open(str(doc_path))
    temp_docx = doc_path.with_suffix('.temp_docx')
    doc.SaveAs(str(temp_docx), FileFormat=16)  # 16 代表 docx
    doc.Close()
    return temp_docx

def pdf_to_md(pdf_path, output_path):
    """提取 PDF 文字和表格"""
    with pdfplumber.open(pdf_path) as pdf:
        content = [f"# Source: {pdf_path.name}"]
        for page in pdf.pages:
            # 提取文本
            text = page.extract_text()
            if text:
                content.append(text)
            # 提取表格并转为 Markdown
            tables = page.extract_tables()
            for table in tables:
                df = pd.DataFrame(table)
                # 过滤掉全空的列
                df = df.dropna(how='all', axis=1)
                content.append(df.to_markdown(index=False))
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n\n'.join(content))

def docx_to_md(docx_path, output_path):
    doc = Document(docx_path)
    md_content = [f"# Source: {docx_path.name}"]
    
    # 按照文档原本顺序处理段落和表格
    for child in doc.element.body:
        if child.tag.endswith('p'):  # 段落
            para = [p for p in doc.paragraphs if p._element == child][0]
            text = para.text.strip()
            if not text: continue
            if para.style.name.startswith('Heading'):
                level = para.style.name.split()[-1]
                md_content.append(f"{'#' * (int(level) if level.isdigit() else 3)} {text}")
            else:
                md_content.append(text)
        elif child.tag.endswith('tbl'):  # 表格
            table = [t for t in doc.tables if t._element == child][0]
            data = [[cell.text.strip().replace('\n', ' ') for cell in row.cells] for row in table.rows]
            if data:
                # 尝试用第一行做表头
                df = pd.DataFrame(data[1:], columns=data[0])
                md_content.append(df.to_markdown(index=False))
                
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n\n'.join(md_content))

def excel_to_md(xlsx_path, output_path):
    """保留你觉得好用的原版 Excel 逻辑"""
    excel_file = pd.ExcelFile(xlsx_path)
    md_content = [f"# Source: {xlsx_path.name}"]
    for sheet_name in excel_file.sheet_names:
        df = pd.read_excel(xlsx_path, sheet_name=sheet_name)
        # 向下填充处理合并单元格，提升 RAG 检索精度
        df = df.ffill(axis=0)
        df = df.dropna(how='all').dropna(axis=1, how='all')
        if not df.empty:
            md_content.append(f"## Sheet: {sheet_name}")
            md_content.append(df.to_markdown(index=False))
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n\n'.join(md_content))

def batch_convert():
    base_dir = Path(__file__).parent.absolute()
    output_dir = base_dir / "converted_md"
    output_dir.mkdir(exist_ok=True)
    
    print(f"--- 正在扫描目录: {base_dir} ---")
    
    count = 0
    for file_path in base_dir.rglob('*'):
        if file_path.name.startswith('~$') or output_dir in file_path.parents:
            continue
            
        ext = file_path.suffix.lower()
        output_file = output_dir / f"{file_path.stem}_{file_path.parent.name}.md"
        
        try:
            if ext == '.docx':
                print(f"处理 Word: {file_path.name}")
                docx_to_md(file_path, output_file)
                count += 1
            elif ext == '.doc':
                print(f"处理旧版 Word: {file_path.name}")
                temp_docx = doc_to_docx(file_path)
                docx_to_md(temp_docx, output_file)
                os.remove(temp_docx)  # 清理临时文件
                count += 1
            elif ext == '.pdf':
                print(f"处理 PDF: {file_path.name}")
                pdf_to_md(file_path, output_file)
                count += 1
            elif ext in ['.xlsx', '.xls']:
                print(f"处理 Excel: {file_path.name}")
                excel_to_md(file_path, output_file)
                count += 1
        except Exception as e:
            print(f"跳过文件 {file_path.name}，原因: {e}")
            
    print(f"--- 处理完成，共转换 {count} 个文件 ---")

if __name__ == "__main__":
    batch_convert()